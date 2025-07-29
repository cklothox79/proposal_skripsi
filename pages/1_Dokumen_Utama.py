import streamlit as st
from docx import Document

st.title("📑 Dokumen Utama")

tab1, tab2, tab3 = st.tabs(["Halaman Judul", "Lembar Persetujuan & Daftar Isi", "Info Tambahan"])

with tab1:
    st.subheader("Halaman Judul")
    try:
        doc = Document("docs/Halaman_Judul.docx")
        for para in doc.paragraphs:
            if para.text.strip():
                st.write(para.text)
    except:
        st.warning("File Halaman_Judul.docx belum ditemukan.")

with tab2:
    st.subheader("Lembar Persetujuan & Daftar Isi")
    try:
        doc = Document("docs/Halaman_Persetujuan_daftar_isi.docx")
        for para in doc.paragraphs:
            if para.text.strip():
                st.write(para.text)
    except:
        st.warning("File ini belum ditemukan.")
